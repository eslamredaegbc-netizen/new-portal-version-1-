from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from monitoring_app.config import APP_NAME, EXPORTS_DIR, FONT_CANDIDATES, OWNER_NAME
from monitoring_app.storage import DatabaseManager
from monitoring_app.utils.text import compact_text, safe_slug


class ReportService:
    def __init__(self, repository: DatabaseManager) -> None:
        self.repository = repository

    def generate_report(
        self,
        *,
        format_name: str,
        report_title: str,
        search_id: int,
        selected_result_ids: List[int],
        executive_summary: str,
    ) -> Dict[str, object]:
        rows = self.repository.export_rows(search_id=search_id, selected_result_ids=selected_result_ids)
        if rows.empty:
            raise ValueError("لا توجد مصادر محددة لإدراجها في تقرير الرصد.")

        summary_text = executive_summary.strip() or self._build_executive_summary(rows, report_title)
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        report_slug = safe_slug(f"{report_title}-{format_name}-{timestamp}")
        path = EXPORTS_DIR / f"{report_slug}.{format_name.lower()}"

        if format_name.lower() == "csv":
            rows.to_csv(path, index=False, encoding="utf-8-sig")
        elif format_name.lower() == "json":
            payload = {
                "generated_at": datetime.utcnow().isoformat(),
                "report_title": report_title,
                "search_id": search_id,
                "executive_summary": summary_text,
                "records": json.loads(rows.fillna("").to_json(orient="records", force_ascii=False, date_format="iso")),
            }
            path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        elif format_name.lower() == "docx":
            self._write_docx(path, rows, report_title, summary_text)
        elif format_name.lower() == "pdf":
            self._write_pdf(path, rows, report_title, summary_text)
        else:
            raise ValueError(f"صيغة غير مدعومة: {format_name}")

        selected_rows = rows.to_dict(orient="records")
        report_id = self.repository.create_report_entry(
            report_name=path.stem,
            report_title=report_title,
            search_id=search_id,
            format_name=format_name.upper(),
            file_path=str(path),
            summary=summary_text,
            selected_rows=selected_rows,
            metadata={
                "search_id": search_id,
                "selected_result_ids": selected_result_ids,
                "generated_at": datetime.utcnow().isoformat(),
            },
        )
        return {
            "report_id": report_id,
            "file_path": str(path),
            "file_name": path.name,
            "format": format_name.upper(),
            "summary": summary_text,
            "selected_results_count": len(selected_rows),
        }

    def _write_docx(self, path: Path, rows, report_title: str, executive_summary: str) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        normal_style.font.size = Pt(11)

        title = document.add_heading(report_title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        subtitle = document.add_paragraph(f"{APP_NAME} | {OWNER_NAME}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        summary_header = document.add_paragraph()
        summary_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        summary_header.add_run("الملخص التنفيذي").bold = True

        summary_text = document.add_paragraph(executive_summary or "لا يوجد ملخص تنفيذي.")
        summary_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for _, row in rows.iterrows():
            heading = document.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            heading.add_run(f"{row['source_name']} | {row['title']}").bold = True

            details = document.add_paragraph()
            details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            details.add_run(
                f"التصنيف: {row['classification']} | اللون: {row['color_label']} | الخطورة: {int(row['risk_score'])}/100"
            )

            legal_summary = document.add_paragraph(compact_text(str(row["legal_summary"]), 900))
            legal_summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            opinion = document.add_paragraph(compact_text(str(row["analyst_opinion"]), 700))
            opinion.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            link = document.add_paragraph(str(row["url"]))
            link.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.save(path)

    def _build_executive_summary(self, rows, report_title: str) -> str:
        top_categories = rows["classification"].fillna("").astype(str).value_counts().head(3)
        top_sources = rows["source_name"].fillna("").astype(str).value_counts().head(3)
        categories_text = "، ".join(f"{name}: {count}" for name, count in top_categories.items()) or "لا توجد فئات بارزة"
        sources_text = "، ".join(f"{name}: {count}" for name, count in top_sources.items()) or "لا توجد مصادر بارزة"
        max_risk = int(rows["risk_score"].fillna(0).max()) if "risk_score" in rows else 0
        red_count = int((rows["color_code"].fillna("") == "red").sum()) if "color_code" in rows else 0
        return compact_text(
            (
                f"يلخص هذا التقرير '{report_title}' اعتمادًا على {len(rows)} مصدرًا مختارًا يدويًا من العضو. "
                f"أعلى درجة خطورة مسجلة هي {max_risk}/100، وعدد النتائج ذات الأولوية الحمراء هو {red_count}. "
                f"أبرز التصنيفات: {categories_text}. أبرز المصادر: {sources_text}."
            ),
            420,
        )

    def _write_pdf(self, path: Path, rows, report_title: str, executive_summary: str) -> None:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        font_name = self._ensure_font()
        document = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            rightMargin=1.2 * cm,
            leftMargin=1.2 * cm,
            topMargin=1.5 * cm,
            bottomMargin=1.5 * cm,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ArabicTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            alignment=2,
            textColor=colors.HexColor("#111111"),
        )
        body_style = ParagraphStyle(
            "ArabicBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=16,
            alignment=2,
            textColor=colors.HexColor("#111111"),
        )
        story = [
            Paragraph(self._shape(report_title), title_style),
            Spacer(1, 0.2 * cm),
            Paragraph(self._shape(f"{APP_NAME} | {OWNER_NAME}"), body_style),
            Spacer(1, 0.15 * cm),
            Paragraph(self._shape("الملخص التنفيذي"), body_style),
            Paragraph(self._shape(executive_summary or "لا يوجد ملخص تنفيذي."), body_style),
            Spacer(1, 0.35 * cm),
        ]

        summary_table = Table(
            [
                [self._shape("المصادر المختارة"), self._shape("أعلى خطورة"), self._shape("الفئات")],
                [
                    self._shape(str(rows.shape[0])),
                    self._shape(str(int(rows["risk_score"].max()))),
                    self._shape("، ".join(rows["classification"].dropna().astype(str).head(4).tolist())),
                ],
            ],
            hAlign="RIGHT",
        )
        summary_table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEDED")),
                    ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D8D8D8")),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("PADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.extend([summary_table, Spacer(1, 0.35 * cm)])

        for _, row in rows.iterrows():
            story.append(Paragraph(self._shape(f"{row['source_name']} | {row['title']}"), body_style))
            story.append(
                Paragraph(
                    self._shape(
                        f"التصنيف: {row['classification']} | اللون: {row['color_label']} | الخطورة: {int(row['risk_score'])}/100"
                    ),
                    body_style,
                )
            )
            story.append(Paragraph(self._shape(compact_text(str(row["legal_summary"]), 700)), body_style))
            story.append(Paragraph(self._shape(compact_text(str(row["analyst_opinion"]), 550)), body_style))
            story.append(Paragraph(self._shape(str(row["url"])), body_style))
            story.append(Spacer(1, 0.25 * cm))

        document.build(story)

    def _ensure_font(self) -> str:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        for candidate in FONT_CANDIDATES:
            if candidate.exists():
                font_name = candidate.stem
                try:
                    if font_name not in pdfmetrics.getRegisteredFontNames():
                        pdfmetrics.registerFont(TTFont(font_name, str(candidate)))
                    return font_name
                except Exception:
                    continue
        return "Helvetica"

    def _shape(self, text: str) -> str:
        try:
            import arabic_reshaper
            from bidi.algorithm import get_display

            return get_display(arabic_reshaper.reshape(text or ""))
        except Exception:
            return text or ""
